import os
import tempfile
import json
import pandas as pd
import numpy as np
import logging
import docx
import PyPDF2
# Add these at the top of your file with other imports
from tqdm import tqdm
import gc
import datetime
import threading
from queue import Queue
import concurrent.futures
import psutil
from typing import Any, Dict, List, Optional, Union, Generator
from dotenv import load_dotenv

# Updated imports for latest LlamaIndex version
from llama_index.core import (
    Document,
    Settings,
    VectorStoreIndex
)
from llama_index.core.node_parser import SimpleNodeParser
from llama_index.core.storage.storage_context import StorageContext
from llama_index.core.retrievers import VectorIndexRetriever
from llama_index.embeddings.openai import OpenAIEmbedding
from llama_index.llms.openai import OpenAI

class EnhancedDocumentRAG:
    def __init__(self, 
                documents_path: str, 
                model_name: str = "gpt-4o-mini",
                embedding_model: str = "text-embedding-ada-002",
                chunk_size: int = 1024,
                chunk_overlap: int = 20,
                batch_size: int = 100,    # New parameter
                max_workers: int = 4):    # New parameter
        """Initialize RAG system with enhanced document processing and analysis"""
        load_dotenv()
        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)

        self.documents_path = documents_path
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.batch_size = batch_size
        self.max_workers = max_workers
        self.processing_queue = Queue()

        # Configure settings using the new Settings approach
        Settings.llm = OpenAI(model=model_name, temperature=0.2)
        Settings.embed_model = OpenAIEmbedding(model=embedding_model)
        Settings.node_parser = SimpleNodeParser.from_defaults(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap
        )

        # Initialize processing
        self.file_analyses = self._comprehensive_document_analysis()
        self.documents = self._load_documents(documents_path)
        self.storage_context = StorageContext.from_defaults()
        self._build_index_efficient()  # Use new efficient index building

    # Step 2.2: Add helper methods for batch processing
    def _batch_generator(self, items: List[Any]) -> Generator:
        """Generate batches from items list"""
        for i in range(0, len(items), self.batch_size):
            yield items[i:i + self.batch_size]
    def process_raw_text(self, text: Union[str, List[str]], text_id: str = None) -> None:
        """
        Process raw text input and add it to the existing index

        Args:
            text: Single string or list of strings to process
            text_id: Optional identifier for the text source
        """
        try:
            # Convert single string to list for uniform processing
            if isinstance(text, str):
                text = [text]

            new_documents = []
            for idx, text_chunk in enumerate(text):
                doc = Document(
                    text=text_chunk,
                    metadata={
                        'source_type': 'raw_text',
                        'text_id': text_id or f'raw_text_{idx}',
                        'chunk_number': idx
                    }
                )
                new_documents.append(doc)

            # Add new documents to existing index
            if hasattr(self, 'index'):
                self.index.refresh_ref_docs(new_documents)
            else:
                self.documents.extend(new_documents)
                self._build_index()

            self.logger.info(f"Successfully processed {len(new_documents)} raw text chunks")

        except Exception as e:
            self.logger.error(f"Error processing raw text: {e}")
            raise

    def process_text_batch(self, 
                          text_batch: List[dict],
                          metadata_fields: List[str] = None) -> None:
        """
        Process a batch of text documents with optional metadata

        Args:
            text_batch: List of dictionaries containing text and metadata
                       Format: [{'text': 'content', 'metadata': {...}}, ...]
            metadata_fields: List of metadata fields to include
        """
        try:
            new_documents = []

            for idx, item in enumerate(text_batch):
                text_content = item.get('text', '')
                if not text_content:
                    continue

                # Process metadata
                metadata = {
                    'source_type': 'batch_text',
                    'batch_id': f'batch_{idx}'
                }

                if metadata_fields and 'metadata' in item:
                    for field in metadata_fields:
                        if field in item['metadata']:
                            metadata[field] = item['metadata'][field]

                doc = Document(
                    text=text_content,
                    metadata=metadata
                )
                new_documents.append(doc)

            # Update index with new documents
            if hasattr(self, 'index'):
                self.index.refresh_ref_docs(new_documents)
            else:
                self.documents.extend(new_documents)
                self._build_index()

            self.logger.info(f"Successfully processed batch of {len(new_documents)} documents")

        except Exception as e:
            self.logger.error(f"Error processing text batch: {e}")
            raise

    def get_text_source_stats(self) -> dict:
        """
        Get statistics about processed text sources
        """
        stats = {
            'total_documents': 0,
            'source_types': {},
            'metadata_fields': set()
        }

        try:
            for doc in self.documents:
                stats['total_documents'] += 1

                # Count source types
                source_type = doc.metadata.get('source_type', 'unknown')
                stats['source_types'][source_type] = stats['source_types'].get(source_type, 0) + 1

                # Track metadata fields
                stats['metadata_fields'].update(doc.metadata.keys())

            # Convert metadata_fields set to list for JSON serialization
            stats['metadata_fields'] = list(stats['metadata_fields'])

            return stats

        except Exception as e:
            self.logger.error(f"Error generating text source statistics: {e}")
            return stats

    def remove_text_source(self, text_id: str) -> bool:
        """
        Remove a specific text source from the index

        Args:
            text_id: Identifier of the text source to remove

        Returns:
            bool: True if successful, False otherwise
        """
        try:
            # Filter out documents with matching text_id
            original_count = len(self.documents)
            self.documents = [doc for doc in self.documents 
                            if doc.metadata.get('text_id') != text_id]

            # Rebuild index if documents were removed
            if len(self.documents) < original_count:
                self._build_index()
                self.logger.info(f"Successfully removed text source: {text_id}")
                return True

            self.logger.info(f"No documents found with text_id: {text_id}")
            return False

        except Exception as e:
            self.logger.error(f"Error removing text source: {e}")
            return False

    def _comprehensive_document_analysis(self) -> Dict[str, Any]:
        """
        Perform comprehensive analysis of documents in the specified directory
        with enhanced statistics tracking
        """
        analysis_results = {
            'file_count': 0,
            'file_types': {},
            'total_size': 0,
            'files': {},
            'dataset_statistics': {}  # New field for detailed dataset stats
        }

        try:
            for filename in os.listdir(self.documents_path):
                file_path = os.path.join(self.documents_path, filename)
                if not os.path.isfile(file_path):
                    continue

                # Get file extension
                file_ext = os.path.splitext(filename)[1].lower()

                # Update file type count
                analysis_results['file_types'][file_ext] = analysis_results['file_types'].get(file_ext, 0) + 1

                # Get file size
                file_size = os.path.getsize(file_path)
                analysis_results['total_size'] += file_size

                # Analyze individual file
                file_info = {
                    'size': file_size,
                    'type': file_ext,
                    'last_modified': os.path.getmtime(file_path),
                    'path': file_path
                }

                # Enhanced CSV analysis
                if file_ext == '.csv':
                    try:
                        df = pd.read_csv(file_path)
                        file_info.update({
                            'row_count': len(df),
                            'column_count': len(df.columns),
                            'columns': list(df.columns),
                            'memory_usage': df.memory_usage(deep=True).sum(),
                            'column_types': df.dtypes.astype(str).to_dict()
                        })
                        # Store detailed dataset statistics
                        analysis_results['dataset_statistics'][filename] = {
                            'rows': len(df),
                            'columns': len(df.columns),
                            'column_names': list(df.columns),
                            'memory_usage_mb': df.memory_usage(deep=True).sum() / (1024 * 1024)
                        }
                    except Exception as e:
                        self.logger.error(f"Error analyzing CSV {filename}: {e}")

                analysis_results['files'][filename] = file_info
                analysis_results['file_count'] += 1

            self.logger.info(f"Analyzed {analysis_results['file_count']} documents")
            return analysis_results

        except Exception as e:
            self.logger.error(f"Error in document analysis: {e}")
            return analysis_results

    def _build_index_efficient(self):
        """Build the vector index with memory-efficient processing"""
        try:
            all_docs = []
            for doc_batch in self._batch_generator(self.documents):
                all_docs.extend(doc_batch)
                gc.collect()

            self.index = VectorStoreIndex.from_documents(
                all_docs,
                storage_context=StorageContext.from_defaults(),
                show_progress=True
            )

        except Exception as e:
            self.logger.error(f"Error building index: {e}")
            raise

    def _load_csv(self, file_path: str) -> List[Document]:
        """Load CSV file with efficient memory handling"""
        documents = []

        # Get total rows for progress bar
        total_rows = sum(1 for _ in pd.read_csv(file_path, chunksize=1))

        # Process in chunks
        with tqdm(total=total_rows, desc="Processing CSV") as pbar:
            for chunk in pd.read_csv(file_path, chunksize=self.batch_size):
                chunk_docs = []
                for idx, row in chunk.iterrows():
                    row_content = "Row Data:\n"
                    for column, value in row.items():
                        row_content += f"{column}: {value}\n"

                    doc = Document(
                        text=row_content,
                        metadata={
                            'filename': os.path.basename(file_path),
                            'row_index': idx,
                            'type': 'csv_row',
                            'processed_date': datetime.datetime.now().isoformat()
                        }
                    )
                    chunk_docs.append(doc)

                documents.extend(chunk_docs)
                pbar.update(len(chunk))
                gc.collect()  # Force garbage collection

        return documents
    def _load_text(self, file_path: str) -> str:
        """Load plain text file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    def _load_docx(self, file_path: str) -> str:
        """Load Microsoft Word document"""
        doc = docx.Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs if para.text])

    def _load_pdf(self, file_path: str) -> str:
        """Load PDF document"""
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            return '\n'.join([page.extract_text() for page in reader.pages])

    def _load_documents(self, documents_path: str) -> List[Document]:
        """
        Load documents with improved handling
        """
        all_documents = []
        supported_extensions = {
            '.csv': self._load_csv,
            '.txt': self._load_text,
            '.docx': self._load_docx,
            '.pdf': self._load_pdf
        }

        for filename in os.listdir(documents_path):
            file_path = os.path.join(documents_path, filename)
            file_ext = os.path.splitext(filename)[1].lower()

            try:
                if file_ext in supported_extensions:
                    load_method = supported_extensions[file_ext]
                    if file_ext == '.csv':
                        # CSV files return a list of documents
                        docs = load_method(file_path)
                        all_documents.extend(docs)
                    else:
                        # Other files return a single document
                        doc_content = load_method(file_path)
                        all_documents.append(Document(
                            text=doc_content,
                            metadata={'filename': filename}
                        ))
            except Exception as e:
                self.logger.error(f"Error loading {filename}: {e}")

        return all_documents

    def retrieve_context(self, query: str, top_k: int = 5) -> List[str]:
        """Memory-efficient context retrieval"""
        try:
            # Create retriever with new API
            retriever = self.index.as_retriever(similarity_top_k=top_k)

            # Retrieve nodes
            retrieved_nodes = retriever.retrieve(query)

            # Process retrieved nodes
            context_chunks = []
            for node in retrieved_nodes:
                chunk = node.text  # or node.get_text() depending on version
                metadata = node.metadata

                if metadata.get('type') == 'csv_row':
                    chunk = f"From {metadata.get('filename', 'unknown file')}, Row {metadata.get('row_index', 'unknown')}: \n{chunk}"

                context_chunks.append(chunk)

            return context_chunks

        except Exception as e:
            self.logger.error(f"Error retrieving context: {e}")
            return []

    def generate_comprehensive_response(
    self, 
        query: str, 
        context_chunks: Optional[List[str]] = None,
        max_chunk_tokens: int = 24000
    ) -> str:
        """
        Generate comprehensive response with improved statistics handling
        """
        try:
            # Check if query is about file statistics
            if any(phrase in query.lower() for phrase in [
                'total', 'rows', 'columns', 'count', 'statistics', 'number of'
            ]):
                # Return file statistics directly from analysis results
                if self.file_analyses and 'dataset_statistics' in self.file_analyses:
                    stats = []
                    for filename, file_stats in self.file_analyses['dataset_statistics'].items():
                        stats.append(f"\nFile: {filename}")
                        stats.append(f"- Number of rows: {file_stats['rows']:,}")
                        stats.append(f"- Number of columns: {file_stats['columns']}")
                        stats.append(f"- Column names: {', '.join(file_stats['column_names'])}")
                        stats.append(f"- Memory usage: {file_stats['memory_usage_mb']:.2f} MB")

                    if stats:
                        return "\n".join(stats)

            # Rest of your existing generate_comprehensive_response code...
            if "display the contents" in query.lower() or "show all data" in query.lower():
                try:
                    df = pd.read_csv(os.path.join(self.documents_path, [f for f in os.listdir(self.documents_path) if f.endswith('.csv')][0]))
                    return f"Dataset contains {len(df)} rows. Here's the full content:\n\n" + df.to_string()
                except Exception as e:
                    return f"Error displaying full dataset: {e}"

            # Original context handling code continues...
            if context_chunks is None:
                context_chunks = self.retrieve_context(query)

            if not context_chunks:
                return "No relevant context found for the query."

            focused_prompt = f"""
            Based on the following context and file statistics, provide a clear and relevant answer to this query: "{query}"

            Focus on:
            1. Use exact numbers from the file statistics when available
            2. Provide clear, quantitative information
            3. Be precise and specific

            Context:
            {' '.join(context_chunks[:3])}  # Limit context chunks

            Answer:
            """

            response = Settings.llm.complete(focused_prompt)
            return str(response.text if hasattr(response, 'text') else response)

        except Exception as e:
            self.logger.error(f"Error generating response: {e}")
            return f"Error generating response: {e}"
    def generate_document_analysis_summary(self) -> str:
        """
        Generate a summary of the document analysis
        """
        analysis = self.file_analyses

        summary = [
            "Document Analysis Summary:",
            f"Total Files: {analysis['file_count']}",
            f"Total Size: {analysis['total_size'] / 1024:.2f} KB",
            "\nFile Types Distribution:"
        ]

        for file_type, count in analysis['file_types'].items():
            summary.append(f"- {file_type}: {count} files")

        summary.append("\nDetailed File Information:")
        for filename, info in analysis['files'].items():
            summary.append(f"\nFile: {filename}")
            summary.append(f"- Size: {info['size'] / 1024:.2f} KB")
            summary.append(f"- Type: {info['type']}")

            if info['type'] == '.csv':
                if 'row_count' in info:
                    summary.append(f"- Rows: {info['row_count']}")
                if 'column_count' in info:
                    summary.append(f"- Columns: {info['column_count']}")

        return "\n".join(summary)
    def optimize_memory(self):
        """Optimize memory usage"""
        if hasattr(self, '_cached_data'):
            del self._cached_data

        gc.collect()

        process = psutil.Process()
        memory_info = process.memory_info()
        self.logger.info(f"Current memory usage: {memory_info.rss / 1024 / 1024:.2f} MB")

# Main execution
# Main execution
# Main execution
if __name__ == "__main__":
    try:
        print("\nEnhanced Document RAG System")
        print("============================")
        print("Enter your input in one of the following formats:")
        print("1. File: <your_file_path>    (e.g., File: C:/documents/)")
        print("2. Text: <your_text>         (e.g., Text: This is my input text)")

        user_input = input("\nYour input: ").strip()
        rag_system = None

        if not user_input:
            print("No input provided. Exiting.")
            exit(1)

        input_parts = user_input.split(':', 1)
        if len(input_parts) != 2:
            print("Invalid input format. Please use 'File: <path>' or 'Text: <content>'")
            exit(1)

        input_type = input_parts[0].lower().strip()
        input_content = input_parts[1].strip()

        if input_type == 'file':
            # Initialize with optimized settings for large datasets
            rag_system = EnhancedDocumentRAG(
                documents_path=input_content,
                chunk_size=1024,
                chunk_overlap=20,
                batch_size=100,  # Adjust based on your system's RAM
                max_workers=4    # Adjust based on your CPU cores
            )
            print(f"\nProcessed documents from: {input_content}")

            # Monitor memory usage
            rag_system.optimize_memory()

        elif input_type == 'text':
            temp_dir = tempfile.mkdtemp()
            try:
                rag_system = EnhancedDocumentRAG(
                    documents_path=temp_dir,
                    chunk_size=1024,
                    chunk_overlap=20,
                    batch_size=100,
                    max_workers=4
                )
                rag_system.process_raw_text(input_content, text_id="user_input_1")
                print("\nProcessed raw text input successfully")

            finally:
                try:
                    os.rmdir(temp_dir)
                except:
                    pass

        # Query processing loop
        while True:
            query = input("\nEnter your question (or 'exit' to quit): ").strip()

            if query.lower() == 'exit':
                break

            if query:
                try:
                    response = rag_system.generate_comprehensive_response(query)
                    print("\n--- Response ---")
                    print(response)

                    # Monitor memory after each query
                    rag_system.optimize_memory()
                except Exception as e:
                    print(f"Error processing query: {e}")
            else:
                print("No question provided. Please try again.")

    except Exception as e:
        print(f"Error: {e}")